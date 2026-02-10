import streamlit as st
import torch
from transformers import pipeline
import nltk
from PyPDF2 import PdfReader
from docx import Document
import io

# Page Configuration
st.set_page_config(page_title="AI Document Humanizer", page_icon="📄")

# --- 1. Load Resources (Cached) ---
@st.cache_resource
def load_model():
    # Fix for NLTK LookupError: Download both 'punkt' and 'punkt_tab'
    nltk.download('punkt')
    nltk.download('punkt_tab')

    # Load model (CPU)
    return pipeline(
        model="Vamsi/T5_Paraphrase_Paws",
        device=-1 
    )

paraphraser = load_model()

# --- 2. Helper Functions ---

def extract_text_from_file(uploaded_file):
    """Extracts text from PDF, DOCX, or TXT files."""
    text = ""
    try:
        if uploaded_file.name.endswith('.pdf'):
            reader = PdfReader(uploaded_file)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        
        elif uploaded_file.name.endswith('.docx'):
            doc = Document(uploaded_file)
            for para in doc.paragraphs:
                text += para.text + "\n"
                
        elif uploaded_file.name.endswith('.txt'):
            text = uploaded_file.read().decode("utf-8")
            
    except Exception as e:
        st.error(f"Error reading file: {e}")
    return text

def create_docx(text):
    """Creates a Word document from text for downloading."""
    doc = Document()
    # Split by newlines to keep some paragraph structure
    for paragraph in text.split('\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph)
    
    # Save to memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def humanize_text(text):
    if not text:
        return ""
    
    sentences = nltk.sent_tokenize(text)
    results = []
    
    # Progress bar
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    total = len(sentences)
    for i, sentence in enumerate(sentences):
        # Update status
        status_text.text(f"Processing sentence {i+1} of {total}...")
        
        # Paraphrase logic
        try:
            prompt = f"paraphrase: {sentence}"
            output = paraphraser(
                prompt,
                max_length=256,
                do_sample=True,
                top_k=50,
                top_p=0.95,
                temperature=1.2,
                num_return_sequences=1
            )
            results.append(output[0]['generated_text'])
        except Exception as e:
            # If a sentence fails, keep the original to prevent crashing
            results.append(sentence)
            
        progress_bar.progress((i + 1) / total)
    
    status_text.empty()
    progress_bar.empty()
    return " ".join(results)

# --- 3. The UI ---
st.title("📄 AI Document Humanizer")
st.markdown("Upload a document (**PDF, DOCX, TXT**) to rewrite it in a human tone.")

# Tabs for different input methods
tab1, tab2 = st.tabs(["📂 Upload File", "✍️ Paste Text"])

input_text = ""

# Handle File Upload
with tab1:
    uploaded_file = st.file_uploader("Choose a file", type=['pdf', 'docx', 'txt'])
    if uploaded_file is not None:
        with st.spinner("Reading file..."):
            input_text = extract_text_from_file(uploaded_file)
            st.info(f"Loaded {len(input_text.split())} words from {uploaded_file.name}")

# Handle Text Paste
with tab2:
    text_area_input = st.text_area("Paste text here", height=200)
    # Priority: If no file uploaded, use pasted text
    if not uploaded_file and text_area_input:
        input_text = text_area_input

# Process Button
if st.button("Humanize Document", type="primary"):
    if input_text:
        with st.spinner("Humanizing content... (This might take a while for long docs)"):
            final_text = humanize_text(input_text)
            
            st.success("Processing Complete!")
            
            # Show Preview
            st.subheader("Preview Result:")
            st.text_area("Result", value=final_text, height=200)
            
            # --- Download Options ---
            col1, col2 = st.columns(2)
            
            # Option 1: Download as Text
            with col1:
                st.download_button(
                    label="⬇️ Download as .txt",
                    data=final_text,
                    file_name="humanized_output.txt",
                    mime="text/plain"
                )
            
            # Option 2: Download as Word Doc
            with col2:
                docx_file = create_docx(final_text)
                st.download_button(
                    label="⬇️ Download as .docx",
                    data=docx_file,
                    file_name="humanized_output.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    else:
        st.warning("Please upload a file or paste text first.")

st.markdown("---")
st.caption("Note: Formatting (bolding, images) from original files may be lost. This tool focuses on text rewriting.")